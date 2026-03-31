"""Generate the complete Claw Code documentation as a Word (.docx) file."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_PATH = Path(__file__).resolve().parent / 'Claw-Code-Documentation.docx'


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Inches(0.3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value


def build_document() -> Document:
    doc = Document()

    # -- Title page --
    title = doc.add_heading('Claw Code', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Documentation technique complete')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Python clean-room rewrite du harness agent Claude Code\ngithub.com/instructkr/claw-code')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_page_break()

    # -- Table des matieres --
    add_heading(doc, 'Table des matieres')
    toc_items = [
        '1. Introduction',
        '2. Installation et utilisation',
        '3. Architecture',
        '4. Concepts cles',
        '5. Tests',
        '6. Mapping TypeScript -> Python',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ===== 1. INTRODUCTION =====
    add_heading(doc, '1. Introduction')

    add_heading(doc, "Qu'est-ce que Claw Code ?", level=2)
    add_para(doc, "Claw Code est une reecriture clean-room en Python du harness agent de Claude Code (Anthropic). "
             "Le projet reproduit l'architecture, les patterns de commandes/outils, le runtime et le query engine "
             "du systeme original, sans copier de code proprietaire.")

    add_heading(doc, 'Contexte', level=2)
    add_para(doc, "Le 31 mars 2026, le code source de Claude Code a ete brievement expose. "
             "Sigrid Jin (@instructkr), power user reconnu par le Wall Street Journal pour avoir consomme "
             "25 milliards de tokens Claude Code, a etudie l'architecture puis reconstruit les patterns en Python.")
    add_para(doc, "Le travail a ete orchestre avec oh-my-codex (OmX), un outil de workflow base sur OpenAI Codex, "
             "en utilisant les modes $team (revue parallele) et $ralph (execution persistante).")

    add_heading(doc, 'Statut du projet', level=2)
    add_table(doc, ['Aspect', 'Detail'], [
        ['Langage', 'Python 3.10+ (zero dependance externe)'],
        ['Port Rust', 'En cours sur la branche dev/rust'],
        ['Parite', 'Miroir structurel, pas encore un remplacement runtime'],
        ['Stars GitHub', '30K+ en quelques heures (record historique)'],
    ])

    add_heading(doc, 'Ce que le projet fait', level=2)
    add_bullet(doc, 'Catalogue 207 commandes et 184 outils du harness original')
    add_bullet(doc, 'Reproduit le pipeline runtime : routing, matching, execution, boucle de tours')
    add_bullet(doc, 'Query engine avec gestion de budget, compaction de transcripts et streaming')
    add_bullet(doc, 'Audit de parite automatique contre l\'archive TypeScript originale')

    add_heading(doc, 'Ce que le projet ne fait PAS', level=2)
    add_bullet(doc, 'Aucun appel API vers un LLM (pas de cle API, pas d\'inference)')
    add_bullet(doc, 'Les commandes/outils sont des shims descriptifs, pas des implementations fonctionnelles')
    add_bullet(doc, 'Les sous-packages sont des stubs structurels (miroir de la hierarchie TS)')

    doc.add_page_break()

    # ===== 2. INSTALLATION =====
    add_heading(doc, '2. Installation et utilisation')

    add_heading(doc, 'Pre-requis', level=2)
    add_bullet(doc, 'Python 3.10 ou superieur')
    add_bullet(doc, 'Aucune dependance externe (bibliotheque standard uniquement)')

    add_heading(doc, 'Cloner le projet', level=2)
    add_code_block(doc, 'git clone https://github.com/instructkr/claw-code.git\ncd claw-code')

    add_heading(doc, 'Commandes principales', level=2)
    add_table(doc, ['Commande', 'Description'], [
        ['python3 -m src.main summary', 'Resume complet du workspace'],
        ['python3 -m src.main manifest', 'Manifeste des modules Python'],
        ['python3 -m src.main parity-audit', 'Audit de parite Python vs TS'],
        ['python3 -m src.main commands --limit 10', 'Lister les commandes mirroir'],
        ['python3 -m src.main tools --limit 10', 'Lister les outils mirroir'],
        ['python3 -m src.main route "prompt" --limit 5', 'Router un prompt'],
        ['python3 -m src.main bootstrap "prompt"', 'Bootstrap une session runtime'],
        ['python3 -m src.main turn-loop "prompt" --max-turns 3', 'Boucle multi-tours'],
        ['python3 -m src.main setup-report', 'Rapport de setup workspace'],
        ['python3 -m src.main command-graph', 'Graphe des commandes'],
        ['python3 -m src.main tool-pool', 'Pool d\'outils assemble'],
        ['python3 -m src.main bootstrap-graph', 'Graphe bootstrap/runtime'],
        ['python3 -m src.main subsystems --limit 16', 'Modules Python du workspace'],
    ])

    add_heading(doc, 'Filtrage avance des outils', level=2)
    add_table(doc, ['Option', 'Effet'], [
        ['--simple-mode', 'Seulement BashTool, FileReadTool, FileEditTool'],
        ['--no-mcp', 'Exclut les outils MCP'],
        ['--deny-prefix mcp', 'Filtre par prefixe'],
        ['--deny-tool BashTool', 'Exclut un outil specifique'],
    ])

    add_heading(doc, 'Autres commandes', level=2)
    add_table(doc, ['Commande', 'Description'], [
        ['show-command <name>', 'Detail d\'une commande'],
        ['show-tool <name>', 'Detail d\'un outil'],
        ['exec-command <name> <prompt>', 'Executer un shim de commande'],
        ['exec-tool <name> <payload>', 'Executer un shim d\'outil'],
        ['flush-transcript <prompt>', 'Persister et flush une session'],
        ['load-session <id>', 'Charger une session sauvegardee'],
        ['remote-mode <target>', 'Mode remote simule'],
        ['ssh-mode <target>', 'Mode SSH simule'],
        ['teleport-mode <target>', 'Mode teleport simule'],
        ['direct-connect-mode <target>', 'Mode direct-connect'],
        ['deep-link-mode <target>', 'Mode deep-link'],
    ])

    add_heading(doc, 'Lancer les tests', level=2)
    add_code_block(doc, 'python3 -m unittest discover -s tests -v')
    add_para(doc, '22 tests couvrent l\'ensemble des fonctionnalites.')

    doc.add_page_break()

    # ===== 3. ARCHITECTURE =====
    add_heading(doc, '3. Architecture')

    add_heading(doc, 'Vue d\'ensemble', level=2)
    add_para(doc, 'Le projet est en Python pur, sans dependance externe. Il utilise dataclasses, json, pathlib, '
             'argparse, uuid et unittest de la bibliotheque standard.')

    add_heading(doc, 'Pipeline core', level=2)
    add_table(doc, ['Etape', 'Module', 'Role'], [
        ['1. Entree CLI', 'main.py', 'Parse argparse, dispatch vers handlers'],
        ['2. Runtime', 'runtime.py', 'Orchestre le cycle de vie : routing, bootstrap, turn loop'],
        ['3. Query Engine', 'query_engine.py', 'Gestion de tours, budget, streaming SSE, persistence'],
        ['4. Registre', 'execution_registry.py', 'Shims executables pour commandes/outils'],
    ])

    add_heading(doc, 'Couche de donnees', level=2)
    add_table(doc, ['Module', 'Role'], [
        ['models.py', 'Dataclasses partages (PortingModule, Subsystem, PermissionDenial, UsageSummary)'],
        ['commands.py', 'Inventaire des commandes depuis commands_snapshot.json'],
        ['tools.py', 'Inventaire des outils depuis tools_snapshot.json'],
        ['session_store.py', 'Persistence JSON des sessions dans .port_sessions/'],
        ['transcript.py', 'Transcript append-only avec compaction et flush'],
    ])

    add_heading(doc, 'Modules de support', level=2)
    add_table(doc, ['Module', 'Role'], [
        ['context.py', 'Contexte du workspace (chemins, compteurs)'],
        ['setup.py', 'Rapport de setup (Python, plateforme, prefetch, deferred init)'],
        ['permissions.py', 'Filtrage d\'outils par deny-list (nom/prefixe)'],
        ['port_manifest.py', 'Introspection du workspace Python'],
        ['parity_audit.py', 'Audit de parite Python vs TypeScript'],
        ['command_graph.py', 'Graphe de segmentation des commandes'],
        ['tool_pool.py', 'Assemblage du pool d\'outils'],
        ['bootstrap_graph.py', 'Graphe des etapes bootstrap/runtime'],
        ['remote_runtime.py', 'Modes simules : remote, SSH, teleport'],
        ['direct_modes.py', 'Modes simules : direct-connect, deep-link'],
        ['history.py', 'Log d\'historique avec entrees horodatees'],
        ['prefetch.py', 'Prefetch side-effects simules'],
        ['deferred_init.py', 'Initialisation differee gated par trust'],
    ])

    add_heading(doc, 'Sous-systemes stubs', level=2)
    add_para(doc, '30+ repertoires sont des packages Python minimaux (__init__.py uniquement) '
             'qui miroir la structure du code TypeScript original :')
    stubs = ('assistant, bootstrap, bridge, buddy, cli, components, constants, coordinator, '
             'entrypoints, hooks, keybindings, memdir, migrations, moreright, native_ts, '
             'outputStyles, plugins, remote, schemas, screens, server, services, skills, '
             'state, types, upstreamproxy, utils, vim, voice')
    add_para(doc, stubs, italic=True)
    add_para(doc, 'Chacun expose MODULE_COUNT et SAMPLE_FILES pour le suivi de parite.')

    doc.add_page_break()

    # ===== 4. CONCEPTS =====
    add_heading(doc, '4. Concepts cles')

    add_heading(doc, 'Mirroring vs Implementation', level=2)
    add_para(doc, 'Claw Code utilise une approche de mirroring : il catalogue et reproduit la structure '
             'du systeme original sans en implementer la logique runtime. Chaque commande et outil est un '
             '"miroir" avec nom, responsabilite, source_hint et status.')

    add_heading(doc, 'Routing de prompts', level=2)
    add_bullet(doc, 'Le prompt est tokenise (split sur espaces, /, -)')
    add_bullet(doc, 'Chaque token est compare par inclusion aux champs de chaque module')
    add_bullet(doc, 'Un score est calcule (nombre de tokens matches)')
    add_bullet(doc, 'Resultats tries par score decroissant')
    add_bullet(doc, 'Au moins une commande et un outil selectionnes si disponibles')

    add_heading(doc, 'Gestion de tours', level=2)
    add_table(doc, ['Parametre', 'Defaut', 'Role'], [
        ['max_turns', '8', 'Nombre maximum de messages'],
        ['max_budget_tokens', '2000', 'Budget total en tokens'],
        ['compact_after_turns', '12', 'Seuil de compaction du transcript'],
        ['structured_output', 'False', 'Mode JSON avec retry'],
        ['structured_retry_limit', '2', 'Tentatives de rendu JSON'],
    ])

    add_heading(doc, 'Streaming SSE', level=2)
    add_table(doc, ['Evenement', 'Contenu'], [
        ['message_start', 'session_id, prompt'],
        ['command_match', 'liste de commandes matchees'],
        ['tool_match', 'liste d\'outils matches'],
        ['permission_denial', 'liste d\'outils refuses'],
        ['message_delta', 'texte de sortie'],
        ['message_stop', 'usage, stop_reason, transcript_size'],
    ])

    add_heading(doc, 'Permission denials', level=2)
    add_bullet(doc, 'Par deny-list : ToolPermissionContext bloque par nom exact ou prefixe')
    add_bullet(doc, 'Par inference : le runtime detecte les outils destructifs (ex: "bash") et les gate')

    add_heading(doc, 'Immutabilite', level=2)
    add_para(doc, 'Tous les dataclasses de donnees utilisent frozen=True. '
             'UsageSummary.add_turn retourne une nouvelle instance, l\'original est inchange. '
             'Seuls QueryEnginePort et PortingBacklog sont mutables (etat de session).')

    add_heading(doc, 'Sessions et persistence', level=2)
    add_bullet(doc, 'Sessions serialisees en JSON dans .port_sessions/')
    add_bullet(doc, 'Cycle : accumulation -> flush_transcript -> persist_session -> load_session')
    add_bullet(doc, 'QueryEnginePort.from_saved_session() reconstruit le moteur depuis une session')

    doc.add_page_break()

    # ===== 5. TESTS =====
    add_heading(doc, '5. Tests')

    add_heading(doc, 'Framework et execution', level=2)
    add_para(doc, 'Les tests utilisent unittest (bibliotheque standard). '
             'Un seul fichier tests/test_porting_workspace.py couvre l\'ensemble du projet.')
    add_code_block(doc, 'python3 -m unittest discover -s tests -v')

    add_heading(doc, 'Liste des 22 tests', level=2)
    add_table(doc, ['Test', 'Verifie'], [
        ['test_manifest_counts_python_files', '>= 20 fichiers Python detectes'],
        ['test_query_engine_summary_mentions_workspace', 'Sections attendues dans le resume'],
        ['test_cli_summary_runs', 'Commande summary OK'],
        ['test_parity_audit_runs', 'Commande parity-audit OK'],
        ['test_root_file_coverage_is_complete_when_local_archive_exists', 'Couverture complete si archive presente'],
        ['test_command_and_tool_snapshots_are_nontrivial', '>= 150 commandes, >= 100 outils'],
        ['test_commands_and_tools_cli_run', 'CLI commands/tools avec filtres'],
        ['test_subsystem_packages_expose_archive_metadata', 'MODULE_COUNT et SAMPLE_FILES exposes'],
        ['test_route_and_show_entry_cli_run', 'Routing et affichage d\'entrees'],
        ['test_bootstrap_cli_runs', 'Bootstrap produit RuntimeSession'],
        ['test_bootstrap_session_tracks_turn_state', 'Matches et usage trackes'],
        ['test_exec_command_and_tool_cli_run', 'Execution de shims'],
        ['test_setup_report_and_registry_filters_run', 'Setup report et filtres'],
        ['test_load_session_cli_runs', 'Persistence et chargement session'],
        ['test_tool_permission_filtering_cli_runs', 'Filtrage deny-prefix'],
        ['test_turn_loop_cli_runs', 'Boucle multi-tours'],
        ['test_remote_mode_clis_run', 'Modes remote/SSH/teleport'],
        ['test_flush_transcript_cli_runs', 'Flush de transcript'],
        ['test_command_graph_and_tool_pool_cli_run', 'Graphe et pool'],
        ['test_setup_report_mentions_deferred_init', 'Deferred init present'],
        ['test_execution_registry_runs', 'Registre d\'execution'],
        ['test_bootstrap_graph_and_direct_modes_run', 'Graphe bootstrap et modes direct'],
    ])

    add_heading(doc, 'Structure des tests', level=2)
    add_para(doc, 'Les tests sont majoritairement des tests d\'integration : '
             'appels CLI via subprocess.run et imports directs de modules. '
             'Les assertions verifient la presence de marqueurs dans stdout.')

    doc.add_page_break()

    # ===== 6. MAPPING =====
    add_heading(doc, '6. Mapping TypeScript -> Python')

    add_heading(doc, 'Fichiers racine', level=2)
    add_table(doc, ['TypeScript', 'Python'], [
        ['QueryEngine.ts', 'QueryEngine.py'],
        ['Task.ts', 'task.py'],
        ['Tool.ts', 'Tool.py'],
        ['commands.ts', 'commands.py'],
        ['context.ts', 'context.py'],
        ['cost-tracker.ts', 'cost_tracker.py'],
        ['costHook.ts', 'costHook.py'],
        ['dialogLaunchers.tsx', 'dialogLaunchers.py'],
        ['history.ts', 'history.py'],
        ['ink.ts', 'ink.py'],
        ['interactiveHelpers.tsx', 'interactiveHelpers.py'],
        ['main.tsx', 'main.py'],
        ['projectOnboardingState.ts', 'projectOnboardingState.py'],
        ['query.ts', 'query.py'],
        ['replLauncher.tsx', 'replLauncher.py'],
        ['setup.ts', 'setup.py'],
        ['tasks.ts', 'tasks.py'],
        ['tools.ts', 'tools.py'],
    ])

    add_heading(doc, 'Repertoires', level=2)
    add_table(doc, ['TypeScript', 'Python'], [
        ['assistant/', 'assistant/'],
        ['bootstrap/', 'bootstrap/'],
        ['bridge/', 'bridge/'],
        ['buddy/', 'buddy/'],
        ['cli/', 'cli/'],
        ['commands/', 'commands.py (agrege)'],
        ['components/', 'components/'],
        ['constants/', 'constants/'],
        ['coordinator/', 'coordinator/'],
        ['entrypoints/', 'entrypoints/'],
        ['hooks/', 'hooks/'],
        ['keybindings/', 'keybindings/'],
        ['memdir/', 'memdir/'],
        ['migrations/', 'migrations/'],
        ['moreright/', 'moreright/'],
        ['native-ts/', 'native_ts/'],
        ['outputStyles/', 'outputStyles/'],
        ['plugins/', 'plugins/'],
        ['remote/', 'remote/'],
        ['schemas/', 'schemas/'],
        ['screens/', 'screens/'],
        ['server/', 'server/'],
        ['services/', 'services/'],
        ['skills/', 'skills/'],
        ['state/', 'state/'],
        ['types/', 'types/'],
        ['upstreamproxy/', 'upstreamproxy/'],
        ['utils/', 'utils/'],
        ['vim/', 'vim/'],
        ['voice/', 'voice/'],
    ])

    add_heading(doc, 'Conventions de nommage', level=2)
    add_bullet(doc, 'Tirets TS -> underscores Python (native-ts -> native_ts)')
    add_bullet(doc, 'Extensions .ts/.tsx -> .py')
    add_bullet(doc, 'Certains repertoires TS agreges en un seul fichier Python')

    return doc


if __name__ == '__main__':
    doc = build_document()
    doc.save(str(OUTPUT_PATH))
    print(f'Document genere : {OUTPUT_PATH}')
