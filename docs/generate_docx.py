"""
Generate a Word document (.docx) for the Claude Code source documentation.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

# Helper functions
def add_code(text, doc=doc):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')

def add_table(headers, rows, doc=doc):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = str(cell)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))

# ==========================================
# PAGE DE TITRE
# ==========================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Claude Code', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Documentation Technique Complete', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph('Analyse du code source (src/)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('~1 884 fichiers TypeScript | ~33 Mo de source')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('CLI officiel d\'Anthropic pour Claude')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

doc.add_page_break()

# ==========================================
# TABLE DES MATIERES
# ==========================================
doc.add_heading('Table des matieres', level=1)
toc_items = [
    '1. Vue d\'ensemble',
    '2. Stack technique',
    '3. Architecture des dossiers',
    '4. Flux d\'initialisation (Bootstrap)',
    '5. Moteur de requetes (QueryEngine & Query)',
    '6. Systeme d\'outils (Tools)',
    '7. Commandes & Skills',
    '8. Services',
    '9. Gestion d\'etat (State)',
    '10. UI & Composants (Ink + React)',
    '11. Multi-agent & Coordination',
    '12. Securite & Permissions',
    '13. Utilitaires & Constantes',
    '14. Patterns architecturaux',
    '15. Diagrammes de flux',
    'Annexe A : Historique et sessions',
    'Annexe B : Configuration et parametres',
    'Annexe C : Types d\'ID',
    'Annexe D : Feature Gates',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ==========================================
# 1. VUE D'ENSEMBLE
# ==========================================
doc.add_heading('1. Vue d\'ensemble', level=1)
doc.add_paragraph(
    'Claude Code est un CLI interactif d\'assistance au developpement logiciel construit par Anthropic. '
    'Il fournit une interface conversationnelle dans le terminal, connectee aux modeles Claude via l\'API '
    'Anthropic, avec un systeme extensible de 60+ outils, 100+ commandes, et un support multi-agent.'
)

doc.add_heading('Fonctionnalites principales', level=2)
add_table(
    ['Categorie', 'Fonctionnalites'],
    [
        ['Outils fichiers', 'Lecture, ecriture, edition, glob, grep, notebooks Jupyter'],
        ['Execution', 'Bash, PowerShell, REPL JavaScript'],
        ['Recherche', 'Web search, web fetch, grep de code'],
        ['IA', 'Streaming API Claude, tool calling, agents autonomes'],
        ['Multi-agent', 'Agents paralleles, equipes (swarms), coordination, worktrees Git'],
        ['MCP', 'Connexion a des serveurs MCP (stdio, SSE, HTTP, WebSocket, SDK)'],
        ['Plugins', 'Systeme de plugins installables (user, project, managed)'],
        ['Skills', 'Commandes personnalisees chargeables depuis le disque'],
        ['Sessions', 'Persistance, resume, historique, compaction de contexte'],
        ['Couts', 'Suivi en temps reel des tokens et couts API par modele'],
        ['Securite', 'Permissions granulaires, classificateur auto, hooks'],
        ['UI', 'Interface terminal React/Ink avec themes, dialogs, diffs'],
    ]
)

doc.add_page_break()

# ==========================================
# 2. STACK TECHNIQUE
# ==========================================
doc.add_heading('2. Stack technique', level=1)

doc.add_heading('Runtime & Langage', level=2)
add_table(
    ['Composant', 'Technologie'],
    [
        ['Langage', 'TypeScript (avec JSX/TSX)'],
        ['Runtime principal', 'Bun'],
        ['Runtime secondaire', 'Node.js 18+'],
        ['Bundler', 'Bun bundler avec bun:bundle (feature gates / dead code elimination)'],
    ]
)

doc.add_heading('Dependances principales', level=2)
add_table(
    ['Package', 'Role'],
    [
        ['@anthropic-ai/sdk', 'Client API Claude (Direct, Bedrock, Vertex, Foundry)'],
        ['@modelcontextprotocol/sdk', 'Protocole MCP (Model Context Protocol)'],
        ['react + ink', 'Rendu UI terminal (React dans le terminal)'],
        ['zod', 'Validation de schemas (inputs outils, configs)'],
        ['chalk', 'Couleurs terminal'],
        ['commander', 'Parsing arguments CLI'],
        ['lodash', 'Utilitaires (memoize, debounce, etc.)'],
        ['@opentelemetry/*', 'Telemetrie et tracing distribue'],
    ]
)

doc.add_page_break()

# ==========================================
# 3. ARCHITECTURE DES DOSSIERS
# ==========================================
doc.add_heading('3. Architecture des dossiers', level=1)

doc.add_heading('Fichiers racine de src/', level=2)
add_table(
    ['Fichier', 'Description'],
    [
        ['main.tsx', 'Point d\'entree, bootstrap de l\'application'],
        ['setup.ts', 'Configuration de session (worktree, hooks, terminal)'],
        ['QueryEngine.ts', 'Boucle principale de requetes (tour par tour)'],
        ['query.ts', 'Orchestration des requetes API'],
        ['Tool.ts', 'Interface abstraite des outils'],
        ['tools.ts', 'Registre et assemblage des outils'],
        ['commands.ts', 'Registre des commandes slash'],
        ['context.ts', 'Construction du contexte (git, memoire, systeme)'],
        ['cost-tracker.ts', 'Suivi des couts et tokens par modele'],
        ['history.ts', 'Historique des sessions (JSONL)'],
        ['Task.ts', 'Gestion des taches en arriere-plan'],
        ['ink.ts', 'Wrapper Ink/React pour le rendu terminal'],
    ]
)

doc.add_heading('Dossiers principaux', level=2)
add_table(
    ['Dossier', 'Role'],
    [
        ['tools/', '43+ implementations d\'outils (Bash, Read, Edit, Agent, Grep, Glob...)'],
        ['commands/', '100+ commandes slash (/commit, /review, /mcp, /skills...)'],
        ['skills/', 'Systeme de skills integrees et chargement depuis le disque'],
        ['services/', 'Couche services (API, MCP, analytics, compaction, plugins)'],
        ['state/', 'Gestion d\'etat (AppStateStore immutable, provider React)'],
        ['components/', '111+ composants React/Ink (UI terminal)'],
        ['hooks/', '75+ hooks React (permissions, notifications, input)'],
        ['coordinator/', 'Mode coordinateur multi-agent'],
        ['assistant/', 'Mode assistant (KAIROS)'],
        ['cli/', 'Interface CLI (REPL, headless, SDK protocol)'],
        ['constants/', 'Constantes globales (prompts, outils, XML tags)'],
        ['types/', 'Definitions de types (Command, Permission, IDs)'],
        ['utils/', '30+ sous-dossiers d\'utilitaires'],
        ['entrypoints/', 'Points d\'entree (init, config, telemetrie)'],
        ['migrations/', 'Migrations de modeles et configurations'],
        ['bridge/', 'Mode bridge (controle a distance)'],
        ['remote/', 'Sessions distantes'],
        ['voice/', 'Mode vocal (feature-gated)'],
        ['plugins/', 'Systeme de plugins'],
        ['vim/', 'Support mode Vim'],
    ]
)

doc.add_page_break()

# ==========================================
# 4. FLUX D'INITIALISATION
# ==========================================
doc.add_heading('4. Flux d\'initialisation (Bootstrap)', level=1)

doc.add_heading('Sequence de demarrage', level=2)
doc.add_paragraph(
    'Le demarrage de Claude Code suit une sequence optimisee avec du prefetching parallele '
    'pour minimiser le temps de lancement.'
)

doc.add_heading('Phase 1 : Prefetch parallele (sans await)', level=3)
add_bullet('startMdmRawRead() - Lecture MDM (plutil/reg query)')
add_bullet('startKeychainPrefetch() - OAuth + cle API legacy')

doc.add_heading('Phase 2 : Chargement des imports (~135ms)', level=3)
doc.add_paragraph('135+ modules charges en parallele avec le prefetch.')

doc.add_heading('Phase 3 : init() [memoize - une seule fois]', level=3)
add_bullet('enableConfigs() - Activation des configurations')
add_bullet('applySafeConfigEnvironmentVariables() - Variables d\'environnement')
add_bullet('applyExtraCACertsFromConfig() - Certificats CA supplementaires')
add_bullet('setupGracefulShutdown() - Arret gracieux')
add_bullet('[async] Logging, OAuth, JetBrains, repository detection')
add_bullet('[async] Remote managed settings, policy limits')
add_bullet('configureGlobalMTLS() - mTLS global')
add_bullet('configureGlobalAgents() - Agents HTTP (proxy)')
add_bullet('preconnectAnthropicApi() - Warmup TCP+TLS')

doc.add_heading('Phase 4 : setup()', level=3)
add_bullet('startUdsMessaging() - Messagerie inter-sessions (Mac/Linux)')
add_bullet('captureTeammateModeSnapshot() - Snapshot pour swarms')
add_bullet('checkAndRestoreTerminalBackup() - Restauration terminal')
add_bullet('setCwd() / captureHooksConfigSnapshot()')
add_bullet('createWorktreeForSession() - Worktree Git si --worktree')
add_bullet('[bg] getCommands(), loadPluginHooks(), registerAttributionHooks()')

doc.add_heading('Phase 5 : Rendu', level=3)
add_bullet('Interactif : launchRepl() - Boucle React/Ink')
add_bullet('Headless : runHeadless() via print.ts - Mode SDK')

doc.add_heading('Phase 6 : Telemetrie post-trust', level=3)
add_bullet('Attente des remote managed settings')
add_bullet('Initialisation OpenTelemetry')

doc.add_heading('Prefixes du prompt systeme', level=2)
add_table(
    ['Contexte', 'Prefixe'],
    [
        ['Interactif (CLI)', 'You are Claude Code, Anthropic\'s official CLI for Claude.'],
        ['SDK (avec append)', 'You are Claude Code, [...] running within the Claude Agent SDK.'],
        ['SDK (sans append)', 'You are a Claude agent, built on Anthropic\'s Claude Agent SDK.'],
        ['Vertex API', 'Toujours le prefixe par defaut'],
    ]
)

doc.add_page_break()

# ==========================================
# 5. MOTEUR DE REQUETES
# ==========================================
doc.add_heading('5. Moteur de requetes', level=1)

doc.add_heading('QueryEngine (src/QueryEngine.ts)', level=2)
doc.add_paragraph(
    'Classe principale qui orchestre la boucle tour par tour entre l\'utilisateur, '
    'le modele et les outils.'
)

doc.add_heading('Configuration du QueryEngine', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['cwd', 'string', 'Repertoire de travail courant'],
        ['tools', 'Tools', 'Outils disponibles'],
        ['commands', 'Command[]', 'Commandes slash'],
        ['mcpClients', 'MCPServerConnection[]', 'Clients MCP connectes'],
        ['agents', 'AgentDefinition[]', 'Definitions d\'agents'],
        ['canUseTool', 'CanUseToolFn', 'Fonction de permission'],
        ['initialMessages', 'Message[]', 'Messages initiaux (resume)'],
        ['readFileCache', 'FileStateCache', 'Cache de lecture fichiers'],
        ['customSystemPrompt', 'string', 'Prompt systeme custom'],
        ['maxTurns', 'number', 'Nombre max de tours'],
        ['maxBudgetUsd', 'number', 'Budget max en USD'],
        ['thinkingConfig', 'ThinkingConfig', 'Config extended thinking'],
    ]
)

doc.add_heading('Boucle de requete', level=2)
doc.add_paragraph('Chaque tour suit ce flux :')
add_bullet('Assemblage du prompt systeme (CLAUDE.md, git status, outils)')
add_bullet('Appel modele : queryModelWithStreaming() avec streaming SSE')
add_bullet('Traitement des evenements : TextDelta, ThinkingDelta, ToolUse')
add_bullet('Execution des outils : runTools() (streaming)')
add_bullet('Verification de compaction : autoCompactIfNeeded()')
add_bullet('Hooks post-sampling')
add_bullet('Suivi du budget (tokens/couts)')
add_bullet('Condition d\'arret : completion, erreur ou max tours')

doc.add_heading('Seuils de compaction automatique', level=2)
add_table(
    ['Constante', 'Valeur', 'Role'],
    [
        ['AUTOCOMPACT_BUFFER_TOKENS', '13 000', 'Buffer pour la compaction auto'],
        ['WARNING_THRESHOLD_BUFFER_TOKENS', '20 000', 'Seuil d\'avertissement'],
        ['ERROR_THRESHOLD_BUFFER_TOKENS', '20 000', 'Seuil d\'erreur'],
        ['MANUAL_COMPACT_BUFFER_TOKENS', '3 000', 'Buffer pour compaction manuelle'],
        ['MAX_CONSECUTIVE_AUTOCOMPACT_FAILURES', '3', 'Disjoncteur (circuit breaker)'],
    ]
)

doc.add_page_break()

# ==========================================
# 6. SYSTEME D'OUTILS
# ==========================================
doc.add_heading('6. Systeme d\'outils (Tools)', level=1)

doc.add_heading('Interface Tool (src/Tool.ts)', level=2)
doc.add_paragraph(
    'Chaque outil implemente cette interface. La fonction buildTool() fournit les valeurs par defaut.'
)

doc.add_heading('Proprietes principales de l\'interface Tool', level=3)
add_table(
    ['Propriete', 'Type', 'Description'],
    [
        ['name', 'string', 'Nom unique de l\'outil'],
        ['aliases', 'string[]', 'Noms alternatifs'],
        ['inputSchema', 'Input', 'Schema d\'entree (Zod)'],
        ['call()', 'async fn', 'Fonction d\'execution principale'],
        ['description()', 'async fn', 'Description pour le modele'],
        ['prompt()', 'async fn', 'Prompt d\'instructions'],
        ['checkPermissions()', 'async fn', 'Verification des permissions'],
        ['isEnabled()', 'boolean', 'Outil actif ?'],
        ['isReadOnly()', 'boolean', 'Lecture seule ?'],
        ['isDestructive()', 'boolean', 'Operation destructive ?'],
        ['isConcurrencySafe()', 'boolean', 'Safe en parallele ?'],
        ['shouldDefer', 'boolean', 'Outil differe (ToolSearch)'],
        ['renderToolUseMessage()', 'ReactNode', 'Rendu UI de l\'invocation'],
    ]
)

doc.add_heading('Catalogue des outils toujours actifs', level=2)
add_table(
    ['Nom', 'Outil', 'Champs d\'entree'],
    [
        ['Bash', 'BashTool', 'command, description?, timeout?, run_in_background?'],
        ['Read', 'FileReadTool', 'file_path, offset?, limit?, pages?'],
        ['Edit', 'FileEditTool', 'file_path, old_string, new_string, replace_all?'],
        ['Write', 'FileWriteTool', 'file_path, content'],
        ['Glob', 'GlobTool', 'pattern, path?'],
        ['Grep', 'GrepTool', 'pattern, path?, glob?, output_mode?, -i?, -n?, -A/-B/-C?, type?, head_limit?, multiline?, offset?'],
        ['Agent', 'AgentTool', 'description, prompt, subagent_type?, model?, run_in_background?, name?, isolation?'],
        ['WebSearch', 'WebSearchTool', 'query (min 2), allowed_domains?, blocked_domains?'],
        ['WebFetch', 'WebFetchTool', 'url, prompt?'],
        ['Skill', 'SkillTool', 'skill, args?'],
        ['TodoWrite', 'TodoWriteTool', 'todos (array: content, activeForm, status)'],
        ['NotebookEdit', 'NotebookEditTool', 'Operations sur cellules Jupyter'],
        ['AskUserQuestion', 'AskUserQuestionTool', 'question'],
        ['EnterPlanMode', 'EnterPlanModeTool', 'Mode planification'],
        ['ExitPlanMode', 'ExitPlanModeV2Tool', 'Sortie du mode plan'],
        ['ToolSearch', 'ToolSearchTool', 'Recherche d\'outils differes'],
    ]
)

doc.add_heading('Outils conditionnels (feature-gated)', level=2)
add_table(
    ['Nom', 'Condition', 'Description'],
    [
        ['TaskCreate/Get/Update/List', 'isTodoV2Enabled', 'Gestion de taches v2'],
        ['EnterWorktree/ExitWorktree', 'isWorktreeModeEnabled', 'Worktrees Git isoles'],
        ['SendMessage', 'Lazy-loaded', 'Messagerie inter-agents'],
        ['CronCreate/Delete/List', 'AGENT_TRIGGERS', 'Taches planifiees'],
        ['RemoteTrigger', 'AGENT_TRIGGERS_REMOTE', 'Declencheurs distants'],
        ['TeamCreate/TeamDelete', 'isAgentSwarmsEnabled', 'Gestion d\'equipes'],
        ['PowerShell', 'isPowerShellToolEnabled', 'Execution PowerShell'],
        ['REPL', 'USER_TYPE=ant', 'REPL JavaScript'],
        ['LSP', 'ENABLE_LSP_TOOL', 'Language Server Protocol'],
        ['WebBrowser', 'WEB_BROWSER_TOOL', 'Automatisation navigateur'],
    ]
)

doc.add_heading('Matrices d\'outils par contexte', level=2)
add_table(
    ['Contexte', 'Outils disponibles'],
    [
        ['CLI principal', 'Tous les outils actifs'],
        ['Sous-agents async', 'Read, Write, Edit, Bash, Glob, Grep, WebSearch, WebFetch, Skill, TodoWrite, NotebookEdit, ToolSearch, Worktree'],
        ['Teammates in-process', 'Outils async + TaskCreate/Get/List/Update + SendMessage + Cron'],
        ['Mode coordinateur', 'Agent, TaskStop, SendMessage, SyntheticOutput uniquement'],
    ]
)

doc.add_page_break()

# ==========================================
# 7. COMMANDES & SKILLS
# ==========================================
doc.add_heading('7. Commandes & Skills', level=1)

doc.add_heading('Types de commandes', level=2)
doc.add_paragraph('Trois types de commandes coexistent :')
add_table(
    ['Type', 'Description', 'Exemple'],
    [
        ['prompt', 'Genere un prompt envoye au modele', '/commit, /review, skills'],
        ['local', 'Executee localement (pas d\'appel API)', '/help, /cost, /clear'],
        ['local-jsx', 'UI React locale interactive', '/init, /config, /mcp'],
    ]
)

doc.add_heading('Commandes principales', level=2)
add_table(
    ['Commande', 'Type', 'Description'],
    [
        ['/help', 'local', 'Aide interactive'],
        ['/init', 'local-jsx', 'Initialisation de projet'],
        ['/config', 'local-jsx', 'Configuration'],
        ['/session', 'local-jsx', 'Gestion de sessions'],
        ['/commit', 'prompt', 'Commit Git assiste'],
        ['/review', 'prompt', 'Revue de code'],
        ['/mcp', 'local-jsx', 'Gestion des serveurs MCP'],
        ['/skills', 'local', 'Listing des skills'],
        ['/plan', 'local', 'Mode planification'],
        ['/cost', 'local', 'Affichage des couts'],
        ['/compact', 'local', 'Compaction manuelle'],
        ['/theme', 'local', 'Selection de theme'],
        ['/model', 'local', 'Selection de modele'],
        ['/permissions', 'local-jsx', 'Gestion des permissions'],
        ['/hooks', 'local', 'Configuration des hooks'],
        ['/vim', 'local', 'Toggle mode Vim'],
        ['/export', 'local', 'Export de la conversation'],
    ]
)

doc.add_heading('Systeme de Skills', level=2)
doc.add_paragraph(
    'Les skills sont des commandes de type "prompt" chargees depuis le disque ou integrees au binaire.'
)

doc.add_heading('Sources de chargement', level=3)
add_bullet('~/.claude/skills/ - Skills utilisateur (globales)')
add_bullet('.claude/skills/ - Skills projet (locales)')
add_bullet('Managed policy settings - Skills d\'entreprise')
add_bullet('Plugin directories - Skills de plugins')
add_bullet('MCP servers - Skills via MCP')
add_bullet('src/skills/bundled/ - Skills integrees au CLI')

doc.add_heading('Skills integrees', level=3)
add_table(
    ['Skill', 'Description'],
    [
        ['batch', 'Execution de prompts en lot'],
        ['debug', 'Utilitaires de debugging'],
        ['loop', 'Planification de taches recurrentes'],
        ['remember', 'Persistance en memoire'],
        ['simplify', 'Revue de simplification de code'],
        ['skillify', 'Creation de nouvelles skills'],
        ['stuck', 'Assistance quand on est bloque'],
        ['verify', 'Boucles de verification'],
        ['keybindings', 'Aide sur les raccourcis clavier'],
        ['claudeApi', 'Exemples d\'utilisation de l\'API Claude'],
    ]
)

doc.add_page_break()

# ==========================================
# 8. SERVICES
# ==========================================
doc.add_heading('8. Services', level=1)

doc.add_heading('8.1 Service API (src/services/api/)', level=2)
doc.add_paragraph(
    'Le service API gere la connexion aux modeles Claude via differents providers.'
)

doc.add_heading('Methodes d\'authentification', level=3)
add_table(
    ['Methode', 'Variable d\'environnement', 'Provider'],
    [
        ['Cle API directe', 'ANTHROPIC_API_KEY', 'API Anthropic'],
        ['AWS Bedrock', 'AWS_ACCESS_KEY_ID + AWS_SECRET_ACCESS_KEY', 'AWS'],
        ['Azure Foundry', 'ANTHROPIC_FOUNDRY_RESOURCE', 'Azure'],
        ['Vertex AI', 'ANTHROPIC_VERTEX_PROJECT_ID', 'Google Cloud'],
        ['OAuth', 'Token en keychain', 'claude.ai / console'],
    ]
)

doc.add_heading('8.2 Service MCP (src/services/mcp/)', level=2)

doc.add_heading('Types de transport', level=3)
add_table(
    ['Transport', 'Description'],
    [
        ['stdio', 'Communication via stdin/stdout d\'un processus'],
        ['sse', 'Server-Sent Events via HTTP'],
        ['sse-ide', 'SSE specifique IDE'],
        ['http', 'HTTP standard (REST)'],
        ['ws', 'WebSocket bidirectionnel'],
        ['sdk', 'SDK interne'],
    ]
)

doc.add_heading('Etats de connexion MCP', level=3)
add_table(
    ['Etat', 'Description'],
    [
        ['connected', 'Connecte avec client, capabilities, cleanup'],
        ['failed', 'Echec avec message d\'erreur'],
        ['needs-auth', 'En attente d\'authentification OAuth'],
        ['pending', 'En cours de connexion (avec tentatives)'],
        ['disabled', 'Desactive par l\'utilisateur'],
    ]
)

doc.add_heading('8.3 Service de compaction (src/services/compact/)', level=2)
add_table(
    ['Module', 'Role'],
    [
        ['compact.ts', 'buildPostCompactMessages() - compression du flux'],
        ['autoCompact.ts', 'autoCompactIfNeeded() - compaction heuristique'],
        ['microCompact.ts', 'microcompactMessages() - compression ciblee'],
        ['snipCompact.ts', 'Compaction par frontieres (feature-gated)'],
        ['grouping.ts', 'Strategies de regroupement de messages'],
    ]
)

doc.add_heading('8.4 Service Analytics (src/services/analytics/)', level=2)
doc.add_paragraph(
    'Backends supportes : Datadog, GrowthBook (feature flags), First-party event logger, OpenTelemetry.'
)
doc.add_paragraph(
    'Types marqueurs pour validation PII : AnalyticsMetadata_I_VERIFIED_THIS_IS_NOT_CODE_OR_FILEPATHS = never'
)

doc.add_heading('8.5 Service Plugins (src/services/plugins/)', level=2)
doc.add_paragraph(
    'Scopes d\'installation valides : user, project, local. '
    'Les plugins peuvent fournir : des outils, des commandes, des hooks, des skills.'
)

doc.add_page_break()

# ==========================================
# 9. GESTION D'ETAT
# ==========================================
doc.add_heading('9. Gestion d\'etat (State)', level=1)

doc.add_heading('Store pattern', level=2)
doc.add_paragraph(
    'Pattern store minimal et immutable avec getState(), setState(updater), subscribe(listener).'
)

doc.add_heading('AppState - champs principaux', level=2)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['settings', 'SettingsJson', 'Configuration persistante'],
        ['mainLoopModel', 'ModelSetting', 'Modele actif'],
        ['tasks', '{[id]: TaskState}', 'Taches en arriere-plan'],
        ['mcp', 'MCP state', 'Clients, outils, ressources MCP'],
        ['plugins', 'Plugin state', 'Plugins actifs/desactives'],
        ['agentDefinitions', 'Agent defs', 'Definitions d\'agents'],
        ['agentNameRegistry', 'Map<string, AgentId>', 'Routage nom vers ID'],
        ['fileHistory', 'File snapshots', 'Historique des fichiers'],
        ['todos', 'Per-agent todos', 'Listes de taches par agent'],
        ['notifications', 'Queue', 'Notifications courantes'],
        ['thinkingEnabled', 'boolean', 'Extended thinking active'],
        ['teamContext', 'Team metadata', 'Equipes, couleurs, panes'],
        ['inbox', 'Message queue', 'File de messages inter-agents'],
        ['speculation', 'Speculation state', 'Suggestions pipelinees'],
        ['denialTracking', 'Denial state', 'Historique des refus'],
    ]
)

doc.add_page_break()

# ==========================================
# 10. UI & COMPOSANTS
# ==========================================
doc.add_heading('10. UI & Composants (Ink + React)', level=1)

doc.add_paragraph(
    'L\'UI est construite avec React + Ink, un framework de rendu React dans le terminal. '
    'Tous les composants sont des composants React standard rendant du texte ANSI.'
)

doc.add_heading('Composants principaux', level=2)
add_table(
    ['Composant', 'Description'],
    [
        ['App.tsx', 'Racine avec providers (theme, state, stats)'],
        ['FullscreenLayout.tsx', 'Layout principal de session (84 Ko)'],
        ['Message.tsx', 'Rendu d\'un message (user/assistant)'],
        ['VirtualMessageList.tsx', 'Liste virtualisee (performance)'],
        ['TextInput.tsx', 'Saisie de texte standard'],
        ['VimTextInput.tsx', 'Saisie en mode Vim'],
        ['StructuredDiff.tsx', 'Affichage de diffs de code'],
        ['ContextVisualization.tsx', 'Visualisation des tokens (76 Ko)'],
        ['StatusLine.tsx', 'Barre de statut en bas'],
        ['GlobalSearchDialog.tsx', 'Recherche dans la conversation'],
        ['Feedback.tsx', 'Soumission de feedback (87 Ko)'],
    ]
)

doc.add_heading('Systeme de design', level=2)
add_bullet('ThemeProvider - Gestion des themes (clair/sombre/custom)')
add_bullet('ThemedBox - Conteneur avec bordures thematiques')
add_bullet('ThemedText - Texte avec couleurs du theme')

doc.add_heading('Modes d\'interface CLI', level=2)
add_table(
    ['Mode', 'Fichier', 'Description'],
    [
        ['Interactif', 'replLauncher.tsx', 'REPL React/Ink complet'],
        ['Headless', 'cli/print.ts', 'Mode non-interactif (SDK, CI)'],
        ['Structure', 'cli/structuredIO.ts', 'Protocole JSON-NDJSON pour SDK'],
        ['Distant', 'cli/remoteIO.ts', 'Sessions via Claude Cloud Relay'],
    ]
)

doc.add_page_break()

# ==========================================
# 11. MULTI-AGENT
# ==========================================
doc.add_heading('11. Multi-agent & Coordination', level=1)

doc.add_heading('Modes multi-agent', level=2)
add_table(
    ['Mode', 'Feature Gate', 'Description'],
    [
        ['Sous-agents', 'Toujours actif', 'Agents lances via AgentTool'],
        ['Coordinateur', 'COORDINATOR_MODE', 'Orchestration multi-workers'],
        ['Equipes (Swarms)', 'isAgentSwarmsEnabled', 'Equipes avec messagerie'],
        ['Teammates', 'Env var', 'Agents en processus avec taches partagees'],
    ]
)

doc.add_heading('Mode coordinateur', level=2)
doc.add_paragraph(
    'Le coordinateur est un orchestrateur qui delegue le travail a des workers. '
    'Il dispose uniquement des outils Agent, SendMessage, TaskStop et SyntheticOutput. '
    'Les workers recoivent un ensemble reduit d\'outils (sans AgentTool pour eviter la recursion).'
)

doc.add_heading('Types de taches', level=2)
add_table(
    ['Type', 'Prefixe ID', 'Description'],
    [
        ['local_bash', 'b', 'Commande shell en arriere-plan'],
        ['local_agent', 'a', 'Agent local'],
        ['remote_agent', 'r', 'Agent distant'],
        ['in_process_teammate', 't', 'Teammate en processus'],
        ['local_workflow', 'w', 'Workflow local'],
        ['monitor_mcp', 'm', 'Moniteur MCP'],
        ['dream', 'd', 'Consolidation memoire'],
    ]
)

doc.add_page_break()

# ==========================================
# 12. SECURITE & PERMISSIONS
# ==========================================
doc.add_heading('12. Securite & Permissions', level=1)

doc.add_heading('Modes de permission', level=2)
add_table(
    ['Mode', 'Comportement'],
    [
        ['default', 'Demande pour chaque outil destructif'],
        ['plan', 'Approbation du plan, puis auto-accept'],
        ['acceptEdits', 'Auto-accept les edits, demande pour bash'],
        ['dontAsk', 'Auto-accept sauf fichiers sensibles'],
        ['bypassPermissions', 'Tout accepter (dangereux)'],
        ['auto', 'Classificateur automatique (ANT-only)'],
        ['bubble', 'Remontee au parent (sous-agents)'],
    ]
)

doc.add_heading('Pipeline de permissions', level=2)
doc.add_paragraph('Chaque demande d\'outil passe par ce pipeline multicouche :')
add_bullet('1. Regles deny/allow configurees (global + projet + policy)')
add_bullet('2. Hooks PreToolUse (scripts configurables)')
add_bullet('3. Classificateur automatique (si mode auto/YOLO)')
add_bullet('4. Dialogue interactif (demande a l\'utilisateur)')
doc.add_paragraph('Chaque couche peut court-circuiter les suivantes.')

doc.add_heading('Sources de decision', level=2)
add_table(
    ['Type', 'Sources possibles'],
    [
        ['Approbation', 'hook (permanent?), user (permanent), classifier'],
        ['Refus', 'hook, user_abort, user_reject (avec feedback?)'],
    ]
)

doc.add_heading('Securite des fichiers de skills', level=2)
add_table(
    ['Protection', 'Mecanisme'],
    [
        ['Anti-symlink', 'O_NOFOLLOW | O_EXCL (Linux), wx (Windows)'],
        ['Owner-only', 'Permissions 0o700 (dirs) et 0o600 (fichiers)'],
        ['Anti-traversal', 'Validation: pas de .. dans les chemins'],
        ['Anti-preemption', 'Nonce par processus dans le repertoire racine'],
    ]
)

doc.add_page_break()

# ==========================================
# 13. UTILITAIRES & CONSTANTES
# ==========================================
doc.add_heading('13. Utilitaires & Constantes', level=1)

doc.add_heading('Sous-dossiers de src/utils/', level=2)
add_table(
    ['Dossier', 'Role'],
    [
        ['bash/', 'Parsing AST de commandes bash (tree-sitter)'],
        ['git/', 'Operations Git (status, diff, config, gitignore)'],
        ['permissions/', 'Systeme de permissions (deny tracking, filesystem)'],
        ['settings/', 'Chargement, validation, fusion de parametres'],
        ['model/', 'Selection de modele, routage de provider'],
        ['hooks/', 'Infrastructure de hooks (post-sampling, session)'],
        ['memory/', 'Gestion de la memoire (CLAUDE.md, attachments)'],
        ['task/', 'Gestion de taches (etat, stockage disque)'],
    ]
)

doc.add_heading('Suivi des couts', level=2)
add_table(
    ['Metrique', 'Description'],
    [
        ['totalCostUSD', 'Cout total de la session en dollars'],
        ['totalAPIDuration', 'Duree totale des appels API'],
        ['totalToolDuration', 'Duree totale d\'execution des outils'],
        ['totalLinesAdded', 'Lignes de code ajoutees'],
        ['totalLinesRemoved', 'Lignes de code supprimees'],
        ['modelUsage', 'Detail par modele (input/output/cache tokens)'],
    ]
)

doc.add_heading('Tags XML pour les messages', level=2)
add_table(
    ['Tag', 'Usage'],
    [
        ['<command-name>', 'Nom de commande slash'],
        ['<bash-input>', 'Commande bash executee'],
        ['<bash-stdout/stderr>', 'Sortie standard/erreur bash'],
        ['<task-notification>', 'Notification de tache (coordinateur)'],
        ['<teammate-message>', 'Message de teammate (swarm)'],
        ['<channel-message>', 'Message de canal'],
        ['<fork-boilerplate>', 'Boilerplate de fork de sous-agent'],
    ]
)

doc.add_page_break()

# ==========================================
# 14. PATTERNS ARCHITECTURAUX
# ==========================================
doc.add_heading('14. Patterns architecturaux', level=1)

doc.add_heading('1. Etat immutable avec store reactif', level=2)
doc.add_paragraph(
    'Toutes les modifications d\'etat passent par setState(prev => newState). '
    'Jamais de mutation directe. Le store supporte les abonnements pour la reactivite.'
)

doc.add_heading('2. Architecture message-driven', level=2)
doc.add_paragraph(
    'Toutes les I/O sont des AsyncIterable<StreamEvent>. Les reponses streamees '
    'sont traitees evenement par evenement (TextDelta, ToolUse, ThinkingDelta, Error...).'
)

doc.add_heading('3. Injection de dependances', level=2)
doc.add_paragraph(
    'QueryDeps permet le mocking pour les tests : callModel, microcompact, autocompact, uuid.'
)

doc.add_heading('4. Feature gates via bundler', level=2)
doc.add_paragraph(
    'feature(\'FEATURE_NAME\') permet la dead code elimination au build. '
    'Le code des features desactivees est supprime du binaire final.'
)

doc.add_heading('5. Memoisation extensive', level=2)
doc.add_paragraph(
    'Les contextes couteux (git status, CLAUDE.md, prompt systeme) sont memoises par session '
    'via lodash memoize() avec invalidation sur changement de dependance.'
)

doc.add_heading('6. Extraction lazy de fichiers', level=2)
doc.add_paragraph(
    'Les skills integrees extraient leurs fichiers de reference sur disque uniquement a la '
    'premiere invocation. La promesse d\'extraction est memorisee pour eviter les conditions de course.'
)

doc.add_heading('7. Compaction adaptive', level=2)
add_table(
    ['Niveau', 'Declencheur', 'Methode'],
    [
        ['Micro', 'Avant chaque requete', 'Compression ciblee de messages specifiques'],
        ['Auto', 'Seuil de tokens depasse', 'Compaction heuristique avec circuit breaker'],
        ['Manuel', 'Commande /compact', 'Compaction complete avec prompt dedie'],
    ]
)

doc.add_heading('8. Permission pipeline multicouche', level=2)
doc.add_paragraph(
    'Regles configurees -> Hooks -> Classificateur -> Dialogue utilisateur. '
    'Chaque couche peut court-circuiter les suivantes. Les decisions sont loguees pour l\'audit.'
)

doc.add_page_break()

# ==========================================
# 15. DIAGRAMMES DE FLUX
# ==========================================
doc.add_heading('15. Diagrammes de flux', level=1)

doc.add_heading('Flux complet d\'un tour de conversation', level=2)
doc.add_paragraph('UTILISATEUR (saisie texte)')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('QueryEngine.submitMessage(prompt)')
doc.add_paragraph('  1. Assemblage prompt systeme (CLAUDE.md, git, outils, skills)')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('API Claude - queryModelWithStreaming()')
doc.add_paragraph('  Streaming SSE : TextDelta | ThinkingDelta | ToolUse')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('Traitement des evenements')
doc.add_paragraph('  TextDelta -> afficher | Thinking -> stocker | ToolUse -> pipeline permission -> executer')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('Post-traitement')
doc.add_paragraph('  Hooks post-sampling | Auto-compaction | Suivi des couts | Mise a jour etat')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('Tour suivant ? (tool_use -> oui, stop -> non)')

doc.add_heading('Flux de permission d\'un outil', level=2)
doc.add_paragraph('Tool.call()')
doc.add_paragraph('    |')
doc.add_paragraph('    v')
doc.add_paragraph('Regles deny/allow ? (global+project+policy) --deny--> REFUSER')
doc.add_paragraph('    | ask/allow')
doc.add_paragraph('    v')
doc.add_paragraph('Hooks PreToolUse ? (scripts configurables) --allow--> AUTORISER')
doc.add_paragraph('    | pas de decision')
doc.add_paragraph('    v')
doc.add_paragraph('Classificateur auto ? (mode auto/YOLO) --allow--> AUTORISER')
doc.add_paragraph('    | pas de decision')
doc.add_paragraph('    v')
doc.add_paragraph('Dialogue utilisateur (terminal interactif) --> Decision user')

doc.add_page_break()

# ==========================================
# ANNEXES
# ==========================================
doc.add_heading('Annexe A : Historique et sessions', level=1)
doc.add_paragraph('Emplacement : ~/.claude/history.jsonl')
doc.add_paragraph('Format : JSON Lines (une entree par ligne)')
doc.add_paragraph('Filtrage par projet (getProjectRoot()) et session (getSessionId())')
doc.add_paragraph('Ordre : session courante en premier, puis tri chronologique inverse')

doc.add_heading('Annexe B : Configuration et parametres', level=1)
add_table(
    ['Fichier', 'Portee', 'Description'],
    [
        ['~/.claude/settings.json', 'Utilisateur', 'Parametres globaux'],
        ['.claude/settings.json', 'Projet', 'Parametres projet'],
        ['.claude/hooks.json', 'Projet', 'Hooks de projet'],
        ['~/.claude/keybindings.json', 'Utilisateur', 'Raccourcis clavier'],
        ['~/.claude/agents/', 'Utilisateur', 'Agents personnalises'],
        ['CLAUDE.md', 'Projet', 'Instructions pour Claude'],
        ['~/.claude/CLAUDE.md', 'Utilisateur', 'Instructions globales'],
    ]
)

doc.add_heading('Annexe C : Types d\'ID', level=1)
doc.add_paragraph('SessionId : string & { __brand: \'SessionId\' } (nominal typing)')
doc.add_paragraph('AgentId : string & { __brand: \'AgentId\' } - Pattern: a(?:.+-)?[0-9a-f]{16}')
doc.add_paragraph('TaskId : [prefixe type][8 chars alphanumeriques] (b=bash, a=agent, r=remote, t=teammate, w=workflow, m=monitor, d=dream)')

doc.add_heading('Annexe D : Feature Gates', level=1)
add_table(
    ['Gate', 'Description'],
    [
        ['COORDINATOR_MODE', 'Orchestration multi-agent'],
        ['KAIROS', 'Mode assistant'],
        ['TRANSCRIPT_CLASSIFIER', 'Classificateur auto de permissions'],
        ['COMMIT_ATTRIBUTION', 'Suivi d\'attribution de commits'],
        ['HISTORY_SNIP', 'Compaction par snipping'],
        ['UDS_INBOX', 'Messagerie inter-sessions'],
        ['CONTEXT_COLLAPSE', 'Optimisation de contexte'],
        ['AGENT_TRIGGERS', 'Taches planifiees (cron)'],
        ['WORKFLOW_SCRIPTS', 'Execution de workflows'],
        ['WEB_BROWSER_TOOL', 'Automatisation navigateur'],
        ['ENABLE_LSP_TOOL', 'Integration LSP'],
        ['PROACTIVE', 'Mode proactif'],
        ['BRIDGE_MODE', 'Controle a distance'],
        ['VOICE_MODE', 'Mode vocal'],
        ['BUDDY', 'Mode buddy'],
        ['FORK_SUBAGENT', 'Fork de sous-agents'],
    ]
)

# ==========================================
# SAVE
# ==========================================
output_path = r'c:\Users\ANAKIN\Downloads\claude_code\docs\DOCUMENTATION.docx'
doc.save(output_path)
print(f'Document Word genere : {output_path}')
